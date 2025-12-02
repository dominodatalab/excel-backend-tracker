# model_registration.py
import os
import hashlib
import base64
import uuid
import time
import shutil
import logging
import tempfile
import json
import re
import pickle
from pathlib import Path
import joblib

import requests
import mlflow
import pandas as pd
from flask import jsonify
from mlflow.pyfunc import PythonModel
from mlflow.models.signature import infer_signature

try:
    from docx import Document
except ImportError:
    Document = None


logger = logging.getLogger(__name__)

DOMINO_DOMAIN = os.environ.get("DOMINO_DOMAIN", "se-demo.domino.tech")
DOMINO_API_KEY = os.environ.get("DOMINO_USER_API_KEY", "")
DOMINO_PROJECT_ID = os.environ.get("DOMINO_PROJECT_ID", "")


def domino_short_id(length: int = 8) -> str:
    """Generate a short ID based on Domino user and project."""
    def short_fallback():
        return base64.urlsafe_b64encode(uuid.uuid4().bytes).decode("utf-8").rstrip("=")[:length]

    user = os.environ.get("DOMINO_USER_NAME") or short_fallback()
    project = os.environ.get("DOMINO_PROJECT_ID") or short_fallback()
    combined = f"{user}/{project}"
    digest = hashlib.sha256(combined.encode()).digest()
    encoded = base64.urlsafe_b64encode(digest).decode("utf-8").rstrip("=")
    return f"{user}_{encoded[:length]}"


EXPERIMENT_NAME = f"external_models_{domino_short_id(4)}"


def _create_pickle_pyfunc():
    """Factory function to create PicklePyFunc class without module dependencies."""
    import pickle
    import pandas as pd
    from mlflow.pyfunc import PythonModel
    
    class PicklePyFunc(PythonModel):
        """MLflow PyFunc wrapper for pickle models."""
        
        def load_context(self, context):
            with open(context.artifacts["model_pkl"], "rb") as f:
                self._model = pickle.load(f)
        
        def predict(self, context, model_input):
            X = model_input if isinstance(model_input, pd.DataFrame) else pd.DataFrame(model_input)
            if hasattr(self._model, "predict_proba"):
                return self._model.predict_proba(X)
            return self._model.predict(X)
    
    return PicklePyFunc()


def send_progress(request_id, step, message, progress_queues, progress=None, file_status=None):
    """Send progress update to the frontend."""
    if request_id in progress_queues:
        progress_queues[request_id].put({
            'step': step,
            'message': message,
            'progress': progress,
            'file_status': file_status
        })


def upload_file_to_project(project_id: str, local_path: str, remote_path: str) -> dict:
    """Upload a file to the head commit of the project repository."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/v4/projects/{project_id}/commits/head/files/{remote_path}"
    headers = {
        "accept": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    
    try:
        with open(local_path, 'rb') as f:
            files = {
                'upfile': (Path(local_path).name, f, 'application/octet-stream')
            }
            logger.info(f"Uploading {local_path} to project {project_id} at {remote_path}")
            response = requests.post(url, headers=headers, files=files, timeout=60)
            response.raise_for_status()
            result = response.json()
            logger.info(f"Successfully uploaded file: {result.get('path')} ({result.get('size')} bytes)")
            return result
    except requests.RequestException as e:
        logger.error(f"Failed to upload {local_path} to project: {e}")
        raise


def attach_report_to_bundle(bundle_id: str, filename: str, commit_key: str) -> dict:
    """Attach a report (HTML or PDF) to a governance bundle."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/governance/v1/bundles/{bundle_id}/attachments"
    headers = {
        "accept": "application/json",
        "Content-Type": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    payload = {
        "identifier": {
            "branch": "master",
            "commit": commit_key,
            "source": "DFS",
            "filename": filename
        },
        "type": "Report"
    }

    try:
        logger.info(f"Attaching report {filename} (commit={commit_key}) to bundle {bundle_id}")
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        result = response.json()
        logger.info(f"Successfully attached report to bundle: {result.get('id')}")
        return result
    except requests.RequestException as e:
        logger.error(f"Failed to attach {filename} to bundle: {e}")
        raise


def save_uploaded_files(files, temp_dir):
    """Save uploaded files to temp directory maintaining structure."""
    saved_files = []
    for file in files:
        filepath = Path(temp_dir) / file.filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        file.save(str(filepath))
        filesize = filepath.stat().st_size
        saved_files.append({
            "path": str(filepath),
            "size_bytes": filesize,
            "size_mb": round(filesize / (1024 * 1024), 2)
        })
    return saved_files


def convert_docx_to_text(docx_path: str) -> str:
    """Convert a .docx file to plain text.

    Args:
        docx_path: Path to the .docx file

    Returns:
        Plain text content of the document
    """
    if Document is None:
        raise ImportError("python-docx library is not installed. Install it with: pip install python-docx")

    try:
        doc = Document(docx_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error converting .docx to text: {e}")
        return f"[Error reading .docx file: {e}]"

def assist_governance_handler(request):
    """Use a Domino gateway LLM to suggest values for governance policy fields based on uploaded files.

    Expects: form fields `policyName`, `policyId`, optional `policy` (JSON string), and file uploads named `files`.
    Returns: JSON { status: 'success', suggestions: { '<label>': '<value>', ... } }
    """
    temp_dir = None

    try:
        # -----------------------------
        # EXTRACT REQUEST DATA
        # -----------------------------
        policy_name = request.form.get('policyName')
        policy_id = request.form.get('policyId')
        policy_json = request.form.get('policy')

        files = request.files.getlist('files')
        temp_dir = tempfile.mkdtemp(prefix=f"assist_{policy_name}_")

        saved_files = save_uploaded_files(files, temp_dir)

        # -----------------------------
        # EXTRACT FILE CONTENTS SAFELY - ONLY PROCESS .DOCX FILES
        # -----------------------------
        file_summaries = {}

        for sf in saved_files:
            path = sf["path"]
            name = Path(path).name
            content = None

            try:
                if name.lower().endswith('.docx'):
                    # Convert .docx to text
                    content = convert_docx_to_text(path)
                else:
                    # Skip non-.docx files
                    content = "[Skipped: Only .docx files are processed]"

            except Exception as e:
                logger.error(f"Error processing file {name}: {e}")
                content = f"[Error processing file: {e}]"

            file_summaries[name] = content

        # -----------------------------
        # EXTRACT ARTIFACT LABELS FROM POLICY JSON
        # -----------------------------
        artifact_labels = []
        range_instructions = []
        if policy_json:
            try:
                pdata = json.loads(policy_json)
                for stage in pdata.get("stages", []):
                    for ev in stage.get("evidenceSet", []):
                        for art in ev.get("artifacts", []):
                            lbl = art.get("details", {}).get("label")
                            if lbl:
                                artifact_labels.append(lbl)
                                # Check if label contains a range pattern like (1-10)
                                range_match = re.search(r'\((\d+)\s*[-–]\s*(\d+)\)', lbl)
                                if range_match:
                                    min_val = range_match.group(1)
                                    max_val = range_match.group(2)
                                    range_instructions.append(
                                        f'- "{lbl}": Return ONLY an integer value between {min_val} and {max_val} (inclusive). No text, no decimals.'
                                    )
            except Exception:
                pass

        # -----------------------------
        # BUILD SYSTEM PROMPT
        # -----------------------------
        range_rules = ""
        if range_instructions:
            range_rules = "\n\nSPECIAL FIELD RULES:\n" + "\n".join(range_instructions)

        system_prompt = (
            "You are a strict JSON-producing governance assistant. "
            "Your ONLY task is to fill evidence values for a Domino governance policy.\n\n"
            "RULES:\n"
            "1. Return ONLY a single JSON object.\n"
            "2. Keys MUST match the artifact labels EXACTLY.\n"
            "3. NO markdown. NO code blocks. NO comments.\n"
            "4. Infer a value whenever probable; use null only if improbable.\n"
            "5. Do not invent keys.\n"
            "6. Be deterministic.\n"
            "7. Be impressive and overly verbose when appropriate.\n"
            f"{range_rules}"
        )

        # -----------------------------
        # BUILD USER PROMPT
        # -----------------------------
        labels_json = json.dumps(artifact_labels, indent=2)
        policy_pretty = policy_json if policy_json else f"Policy ID: {policy_id}"

        user_prompt_parts = [
            "POLICY JSON:",
            policy_pretty,
            "\nARTIFACT LABELS TO FILL:",
            labels_json,
            "\nFILES:",
        ]

        for fname, summary in file_summaries.items():
            user_prompt_parts.append(f"\n--- FILE: {fname} ---\n{summary}")

        user_prompt_parts.append(
            "\nReturn ONLY a JSON object.\n"
            "{\n"
            '  "Model Name": "...",\n'
            '  "Model Description": "..."\n'
            "}"
        )

        full_user_prompt = "\n".join(user_prompt_parts)

        # -----------------------------
        # CALL LLM
        # -----------------------------
        from mlflow.deployments import get_deploy_client
        client = get_deploy_client(os.environ["DOMINO_MLFLOW_DEPLOYMENTS"])
        endpoint = os.environ.get("DOMINO_GATEWAY_LLM_ENDPOINT", "fsi-chatbot")

        response = client.predict(
            endpoint=endpoint,
            inputs={
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": full_user_prompt},
                ],
                "temperature": 0,
                "top_p": 0.1,
                "seed": 42,
                "max_tokens": 1000
            },
        )

        # -----------------------------
        # EXTRACT RAW STRING RESPONSE
        # -----------------------------
        suggestions_raw = None
        if isinstance(response, dict):
            for k in ("predictions", "outputs", "result", "text"):
                if k in response:
                    suggestions_raw = response[k]
                    break
            if suggestions_raw is None:
                suggestions_raw = json.dumps(response)
        else:
            suggestions_raw = str(response)

        if isinstance(suggestions_raw, (dict, list)):
            suggestions_str = json.dumps(suggestions_raw)
        else:
            suggestions_str = str(suggestions_raw)
        print('9'*80)
        print(suggestions_str)
        # -----------------------------
        # PARSE JSON SAFELY
        # -----------------------------
        try:
            suggestions = json.loads(suggestions_str)
        except Exception:
            # fallback: extract the first `{...}`
            m = re.search(r"\{[\s\S]*\}", suggestions_str)
            if m:
                try:
                    suggestions = json.loads(m.group(0))
                except Exception:
                    suggestions = {}
            else:
                suggestions = {}

        # clean dict
        clean = {}
        if isinstance(suggestions, dict):
            for k, v in suggestions.items():
                clean[str(k).strip()] = v

        return jsonify({"status": "success", "suggestions": clean}), 200

    except Exception as e:
        logger.error(f"assist_governance_handler failed: {e}", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500

    finally:
        if temp_dir and Path(temp_dir).exists():
            shutil.rmtree(temp_dir)


def update_model_description(model_name: str, description: str) -> dict:
    """Update model description via Domino API and return full response."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/registeredmodels/v1/{model_name}"
    headers = {
        "accept": "application/json",
        "Content-Type": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    payload = {
        "description": description,
        "discoverable": True
    }
    
    try:
        logger.info(f"Updating model description for {model_name}")
        response = requests.patch(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        model_data = response.json()
        logger.info(f"Successfully updated model description for {model_name}")
        return model_data
    except requests.RequestException as e:
        logger.error(f"Failed to update model description: {e}")
        raise


def normalize_label(label: str) -> str:
    x = label.lower().strip()
    x = re.sub(r"\(\s*\d+\s*[-–]\s*\d+\s*\)", "", x)
    x = re.sub(r"[^\w\s-]", "", x)
    x = re.sub(r"\s+", "_", x)

    return x


def get_policy_details(policy_id: str) -> dict:
    """Get policy details including classification artifact map."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/governance/v1/policies/{policy_id}"
    headers = {
        "accept": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    
    try:
        logger.info(f"Getting policy details for policy ID: {policy_id}")
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        policy_data = response.json()
        logger.info(f"Successfully retrieved policy: {policy_data.get('name')}")
        
        tuples = []
        policy_id_from_response = policy_data.get("id")
        stages = policy_data.get("stages", [])
        
        for stage in stages:
            for evidence in stage.get("evidenceSet", []):
                evidence_id = evidence.get("id")
                for artifact in evidence.get("artifacts", []):
                    artifact_id = artifact.get("id")
                    label = artifact.get("details", {}).get("label")
                    input_type = artifact.get("details", {}).get("type")
                    if policy_id_from_response and evidence_id and artifact_id and label and input_type:
                        tuples.append((policy_id_from_response, evidence_id, artifact_id, label, input_type))
        
        unique_tuples = list(dict.fromkeys(tuples))
        
        return policy_data
    except requests.RequestException as e:
        logger.error(f"Failed to get policy details: {e}")
        raise


def create_bundle(model_name: str, model_version: int, policy_id: str) -> dict:
    """Create a bundle for the registered model."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/governance/v1/bundles"
    headers = {
        "accept": "application/json",
        "Content-Type": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    
    bundle_name = f"{model_name}_v{model_version}"
    payload = {
        "attachments": [
            {
                "identifier": {
                    "name": model_name,
                    "version": model_version
                },
                "type": "ModelVersion"
            }
        ],
        "name": bundle_name,
        "policyId": policy_id,
        "projectId": DOMINO_PROJECT_ID
    }
    
    try:
        logger.info(f"Creating bundle: {bundle_name}")
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        bundle_data = response.json()
        logger.info(f"Successfully created bundle: {bundle_data.get('id')}")
        return bundle_data
    except requests.RequestException as e:
        logger.error(f"Failed to create bundle: {e}")
        raise


def submit_artifacts_to_policy(bundle_id: str, policy_id: str, matched_artifacts: list) -> dict:
    """Submit all artifacts to policy in a single batch call."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/governance/v1/rpc/submit-result-to-policy"
    headers = {
        "accept": "application/json",
        "Content-Type": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }
    
    evidence_groups = {}
    for artifact in matched_artifacts:
        if artifact['value'] is not None:
            evidence_id = artifact['evidence_id']
            if evidence_id not in evidence_groups:
                evidence_groups[evidence_id] = []
            evidence_groups[evidence_id].append(artifact)
    
    logger.info(f"Submitting {len(matched_artifacts)} artifacts across {len(evidence_groups)} evidence sets")
    
    results = []
    for evidence_id, artifacts in evidence_groups.items():
        content = {}
        for artifact in artifacts:
            artifact_id = artifact['artifact_id']
            value = artifact['value']
            
            if artifact['input_type'] == 'radio':
                if isinstance(value, bool):
                    content[artifact_id] = "Yes" if value else "No"
                else:
                    content[artifact_id] = str(value)
            else:
                content[artifact_id] = str(value)
        
        payload = {
            "bundleId": bundle_id,
            "content": content,
            "evidenceId": evidence_id,
            "policyId": policy_id
        }
        
        try:
            logger.info(f"Submitting evidence group {evidence_id} with {len(content)} artifacts")
            response = requests.post(url, json=payload, headers=headers, timeout=30)
            response.raise_for_status()
            result_data = response.json()
            results.append(result_data)
            logger.info(f"Successfully submitted evidence group {evidence_id}")
        except requests.RequestException as e:
            logger.error(f"Failed to submit evidence group {evidence_id}: {e}")
            raise
    
    return results[-1] if results else {}


def create_bundle_simple(bundle_name: str, policy_id: str) -> dict:
    """Create a simple governance bundle without model version attachment."""
    domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
    url = f"https://{domain}/api/governance/v1/bundles"
    headers = {
        "accept": "application/json",
        "Content-Type": "application/json",
        "X-Domino-Api-Key": DOMINO_API_KEY
    }

    payload = {
        "attachments": [],
        "name": bundle_name,
        "policyId": policy_id,
        "projectId": DOMINO_PROJECT_ID
    }

    try:
        logger.info(f"Creating bundle: {bundle_name}")
        response = requests.post(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        bundle_data = response.json()
        logger.info(f"Successfully created bundle: {bundle_data.get('id')}")
        return bundle_data
    except requests.RequestException as e:
        logger.error(f"Failed to create bundle: {e}")
        raise


def register_model_handler(request, progress_queues):
    """Handle governance data submission without model registration or security scanning."""
    logger.info("=" * 80)
    logger.info("SUBMIT GOVERNANCE DATA - Request Received")
    logger.info("=" * 80)

    temp_dir = None
    request_id = request.form.get("requestId", str(uuid.uuid4()))

    try:
        # Get policy information
        policy_name = request.form.get("policyName")
        policy_id = request.form.get("policyId")

        # Parse dynamic fields from JSON
        dynamic_fields_json = request.form.get("dynamicFields", "{}")
        dynamic_fields = json.loads(dynamic_fields_json)

        logger.info(f"Policy Name: {policy_name}")
        logger.info(f"Policy ID: {policy_id}")
        logger.info(f"Dynamic Fields: {dynamic_fields}")

        if not policy_id or not policy_name:
            return jsonify({"status": "error", "message": "Policy selection is required"}), 400

        send_progress(request_id, 'policy', 'Retrieving policy details...', progress_queues, progress=10)
        policy_data = get_policy_details(policy_id)

        files = request.files.getlist('files')
        temp_dir = tempfile.mkdtemp(prefix=f"governance_{policy_name}_")
        logger.info(f"Created temp directory: {temp_dir}")

        send_progress(request_id, 'upload', 'Saving uploaded files...', progress_queues, progress=20)

        saved_files = save_uploaded_files(files, temp_dir)
        logger.info(f"Saved {len(saved_files)} files to temp directory")

        # Extract text from .docx files
        send_progress(request_id, 'extract', 'Extracting text from .docx files...', progress_queues, progress=30)
        file_contents = {}
        for saved_file in saved_files:
            filepath = saved_file.get('path', '')
            filename = Path(filepath).name

            if filename.lower().endswith('.docx'):
                text_content = convert_docx_to_text(filepath)
                file_contents[filename] = text_content
                logger.info(f"Extracted text from {filename}: {len(text_content)} characters")

        # Generate unique bundle name from model_name plus short random suffix
        base_name = dynamic_fields.get('model_name', 'governance_bundle')
        bundle_name = f"{base_name}_{uuid.uuid4().hex[:4]}"

        send_progress(request_id, 'bundle', 'Creating governance bundle...', progress_queues, progress=50)
        bundle_data = create_bundle_simple(bundle_name, policy_id)
        bundle_id = bundle_data.get("id", "")
        project_owner = bundle_data.get("projectOwner", "")
        project_name = bundle_data.get("projectName", "")
        project_id = bundle_data.get("projectId", "")
        stage = bundle_data.get("stage", "").lower().replace(" ", "-")

        # Build matched artifacts from policy and dynamic fields
        stages = policy_data.get("stages", [])

        matched_artifacts = []
        seen_keys = set()

        for stage in stages:
            for evidence in stage.get("evidenceSet", []):
                evidence_id = evidence.get("id")
                for artifact in evidence.get("artifacts", []):
                    artifact_id = artifact.get("id")
                    label = artifact.get("details", {}).get("label")
                    input_type = artifact.get("details", {}).get("type")

                    if policy_id and evidence_id and artifact_id and label and input_type:
                        unique_key = (policy_id, evidence_id, artifact_id, label, input_type)

                        if unique_key not in seen_keys:
                            seen_keys.add(unique_key)
                            # Try to find matching value in dynamic fields
                            # IMPORTANT: Initialize value to None for each artifact
                            value = None

                            normalized_key = normalize_label(label)
                            for field_key, field_value in dynamic_fields.items():
                                if normalize_label(field_key) == normalized_key:
                                    value = field_value
                                    break

                            matched_artifacts.append({
                                'bundle_id': bundle_id,
                                'policy_id': policy_id,
                                'evidence_id': evidence_id,
                                'artifact_id': artifact_id,
                                'label': label,
                                'input_type': input_type,
                                'value': value
                            })

        logger.info("\nMatched Policy Artifacts with Evidence Variables:")
        for artifact in matched_artifacts:
            if artifact['value'] is not None:
                logger.info(f"  {artifact['label']}: {artifact['value']}")

        send_progress(request_id, 'evidence', 'Submitting evidence to policy...', progress_queues, progress=70)
        policy_submission_result = submit_artifacts_to_policy(bundle_id, policy_id, matched_artifacts)
        logger.info(f"Successfully submitted {len([a for a in matched_artifacts if a['value'] is not None])} artifacts to policy")

        send_progress(request_id, 'complete', 'Submission complete!', progress_queues, progress=100)
        send_progress(request_id, 'done', '', progress_queues, progress=100)

        if temp_dir and Path(temp_dir).exists():
            shutil.rmtree(temp_dir)
            logger.info(f"Cleaned up temp directory: {temp_dir}")

        logger.info("=" * 80)

        domain = DOMINO_DOMAIN.removeprefix("https://").removeprefix("http://")
        bundle_url = f"https://{domain}/u/{project_owner}/{project_name}/governance/bundle/{bundle_id}/policy/{policy_id}/evidence/stage/{stage}"

        response_data = {
            "status": "success",
            "message": "Governance data submitted successfully",
            "data": {
                "file_count": len(saved_files),
                "bundle_id": bundle_id,
                "bundle_name": bundle_data.get("name"),
                "bundle_url": bundle_url,
                "policy_name": policy_name,
                "policy_id": policy_id,
                "project_id": DOMINO_PROJECT_ID,
                "project_name": project_name,
                "artifacts_submitted": len([a for a in matched_artifacts if a['value'] is not None])
            }
        }

        return jsonify(response_data), 200

    except Exception as e:
        logger.error(f"Error submitting governance data: {e}", exc_info=True)
        send_progress(request_id, 'error', f'Error: {str(e)}', progress_queues, progress=0)
        send_progress(request_id, 'done', '', progress_queues, progress=0)

        if temp_dir and Path(temp_dir).exists():
            shutil.rmtree(temp_dir)

        return jsonify({
            "status": "error",
            "message": f"Failed to submit governance data: {str(e)}"
        }), 500